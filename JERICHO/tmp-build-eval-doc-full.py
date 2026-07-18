#!/usr/bin/env python3
"""Build the full-horizon Word evaluation doc from the 1072-block extract."""
import json
from collections import defaultdict
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path("/Users/jamesdotson/vscode/JERICHO/JERICHO")
SRC = ROOT / "tmp-schedule-extract-full.json"
OUT = ROOT / "tmp-jericho-evaluation-full.docx"

bundle = json.loads(SRC.read_text())
meta = bundle["meta"]
mp = bundle["masterPlan"]
summary = bundle["summary"]
lanes = bundle["lanes"]
milestones = bundle["milestones"]
cycle_blocks = bundle["cycleBlocks"]
fh_blocks = bundle["fullHorizonBlocks"]
cycle = bundle["cycleSummary"]

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)


def shade(cell, color):
    pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color)
    pr.append(shd)


def add_header(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    return h


def add_kv(label, value):
    p = doc.add_paragraph()
    p.add_run(f"{label}: ").bold = True
    p.add_run("" if value is None else str(value))


def make_table(headers):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, label in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = label
        shade(cell, "1F2A44")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    return t


# ===================== COVER =====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Jericho — Full 5-Year Schedule Evaluation")
tr.font.size = Pt(22)
tr.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run(mp.get("title", "(untitled plan)"))
sr.font.size = Pt(14)
sr.italic = True

stamp = doc.add_paragraph()
stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
stamp.add_run(
    f"View date: {meta.get('viewDate','?')}  •  "
    f"Lifecycle: {meta.get('goalLifecycleState','?')}  •  "
    f"Horizon: {meta.get('range',{}).get('startDayKey','?')} → {meta.get('range',{}).get('endDayKey','?')}  •  "
    f"{len(fh_blocks)} blocks"
).italic = True

extracted = meta.get("extractedAtISO", "")
ex = doc.add_paragraph()
ex.alignment = WD_ALIGN_PARAGRAPH.CENTER
ex.add_run(f"Extracted: {extracted}  •  agenda: {meta.get('agendaVersionId','?')[:50]}").font.size = Pt(9)

doc.add_paragraph()

# ===================== 1. PLAN OVERVIEW =====================
add_header("1. Plan overview", level=1)

add_kv("Plan ID", mp.get("id"))
add_kv("Status", mp.get("status"))
add_kv("Horizon", f"{mp.get('horizonStart')} → {mp.get('horizonEnd')}  ({mp.get('declaredHorizonMonths')} months)")
add_kv("Full-horizon end day key", mp.get("fullHorizonEndDayKey"))
add_kv("Official start", mp.get("officialStartDate"))
add_kv("Schedule applied", mp.get("scheduleAppliedDate"))
add_kv("Controllability class", mp.get("controllabilityClass"))
add_kv("Terminal target class", mp.get("terminalTargetClass"))
add_kv("Goal architecture", mp.get("goalArchitecture"))
add_kv("Execution model", mp.get("executionModel"))
add_kv("Primary lane", mp.get("primaryLane"))

add_header("North-star outcome", level=2)
doc.add_paragraph(mp.get("northStarOutcome") or "(none)")
add_header("Core mission", level=2)
doc.add_paragraph(mp.get("coreMission") or "(none)")
add_header("Outcome target", level=2)
doc.add_paragraph(mp.get("outcomeTarget") or "(none)")
add_header("Success standard", level=2)
doc.add_paragraph(mp.get("successStandard") or "(none)")
add_header("Plan summary", level=2)
doc.add_paragraph(mp.get("masterPlanSummary") or "(none)")

add_header("Anchors", level=2)
for a in mp.get("anchors") or []:
    if isinstance(a, dict):
        doc.add_paragraph(f"{a.get('date','?')} — {a.get('label','(no label)')}", style="List Bullet")
    else:
        doc.add_paragraph(str(a), style="List Bullet")

doc.add_page_break()

# ===================== 2. LANES =====================
add_header("2. Lanes", level=1)
doc.add_paragraph(f"{len(lanes)} parallel workstreams.")

t = make_table(["Lane", "Domain", "Role", "Activation", "Stage", "Milestones"])
for lane_id, lane in lanes.items():
    row = t.add_row().cells
    row[0].text = lane.get("title") or "(untitled)"
    row[1].text = lane.get("domain") or ""
    row[2].text = lane.get("role") or ""
    row[3].text = lane.get("activationState") or ""
    row[4].text = lane.get("assessedStage") or ""
    row[5].text = str(lane.get("milestoneCount") or 0)

doc.add_page_break()

# ===================== 3. MILESTONES =====================
add_header("3. Milestones", level=1)
doc.add_paragraph(f"{len(milestones)} milestones across {len(lanes)} lanes, sorted by target date.")

t = make_table(["Target", "Lane", "Milestone", "Type", "Flex"])
for m in milestones:
    row = t.add_row().cells
    row[0].text = m.get("targetDate") or ""
    row[1].text = (m.get("laneTitle") or "")[:30]
    row[2].text = m.get("title") or ""
    row[3].text = m.get("milestoneType") or ""
    row[4].text = m.get("flex") or ""

doc.add_page_break()

# ===================== 4. SCHEDULE — STRATEGIC SHAPE =====================
add_header("4. Full-horizon schedule — strategic shape", level=1)

total_hours = sum((b.get("durationMinutes") or 0) for b in fh_blocks) / 60
doc.add_paragraph(
    f"{len(fh_blocks)} blocks materialize the full 5-year horizon, totaling "
    f"{total_hours:.0f} hours of planned work. "
    f"Below: how those blocks distribute by phase, year, quarter, lane, and type."
)

add_header("4.1 By phase", level=2)
t = make_table(["Phase", "Block count", "Share"])
for ph, n in sorted((summary.get("byPhase") or {}).items()):
    row = t.add_row().cells
    row[0].text = ph
    row[1].text = str(n)
    row[2].text = f"{(n / max(1, len(fh_blocks))) * 100:.0f}%"

add_header("4.2 By year", level=2)
t = make_table(["Year", "Block count", "Approx hours"])
year_hours = defaultdict(float)
for b in fh_blocks:
    yr = (b.get("dayKey") or "")[:4]
    if yr:
        year_hours[yr] += (b.get("durationMinutes") or 0) / 60
for yr, n in sorted((summary.get("byYear") or {}).items()):
    row = t.add_row().cells
    row[0].text = yr
    row[1].text = str(n)
    row[2].text = f"{year_hours.get(yr, 0):.0f}"

add_header("4.3 By quarter", level=2)
t = make_table(["Quarter", "Block count"])
for q, n in sorted((summary.get("byQuarter") or {}).items()):
    row = t.add_row().cells
    row[0].text = q
    row[1].text = str(n)

add_header("4.4 By lane", level=2)
t = make_table(["Lane", "Block count", "Share"])
lane_name_by_id = {lid: lane.get("title") for lid, lane in lanes.items()}
for lid, n in sorted((summary.get("byLane") or {}).items(), key=lambda kv: -kv[1]):
    row = t.add_row().cells
    row[0].text = (lane_name_by_id.get(lid) or lid)[:50]
    row[1].text = str(n)
    row[2].text = f"{(n / max(1, len(fh_blocks))) * 100:.0f}%"

add_header("4.5 By block type", level=2)
t = make_table(["Type", "Block count", "Share"])
for tp, n in sorted((summary.get("byBlockType") or {}).items(), key=lambda kv: -kv[1]):
    row = t.add_row().cells
    row[0].text = tp
    row[1].text = str(n)
    row[2].text = f"{(n / max(1, len(fh_blocks))) * 100:.0f}%"

doc.add_page_break()

# ===================== 5. SCHEDULE — YEAR-BY-YEAR DETAIL =====================
add_header("5. Year-by-year schedule detail", level=1)
doc.add_paragraph(
    "Every block, grouped by year and lane. Within a lane, sorted by date. "
    "Use this section to audit whether the right work is planned at the right time."
)

by_year = defaultdict(lambda: defaultdict(list))
for b in fh_blocks:
    yr = (b.get("dayKey") or "")[:4] or "?"
    lane_title = b.get("laneTitle") or "(cross-lane)"
    by_year[yr][lane_title].append(b)

for yr in sorted(by_year.keys()):
    yr_total = sum((b.get("durationMinutes") or 0) for blocks in by_year[yr].values() for b in blocks) / 60
    add_header(f"{yr}  —  {sum(len(blocks) for blocks in by_year[yr].values())} blocks, {yr_total:.0f} hours", level=2)
    for lane_title in sorted(by_year[yr].keys()):
        blocks = sorted(by_year[yr][lane_title], key=lambda b: (b.get("dayKey") or "", b.get("title") or ""))
        add_header(f"{lane_title}  ({len(blocks)} blocks)", level=3)
        t = make_table(["Date", "Phase", "Type", "Title", "Min"])
        for b in blocks:
            row = t.add_row().cells
            row[0].text = b.get("dayKey") or ""
            row[1].text = b.get("phaseLabel") or ""
            row[2].text = b.get("blockType") or ""
            row[3].text = (b.get("title") or "")[:140]
            row[4].text = str(b.get("durationMinutes") or "")
    doc.add_page_break()

# ===================== 6. ACTIVE CYCLE (executable now) =====================
add_header("6. Active cycle — executable now", level=1)
doc.add_paragraph(
    f"The active cycle materializes {len(cycle_blocks)} blocks into the executable layer. "
    "These are the blocks that drive the current Today/Week/Month views."
)

t = make_table(["Start", "Min", "Lane", "Title"])
for b in cycle_blocks:
    row = t.add_row().cells
    row[0].text = b.get("startISO") or b.get("dayKey") or ""
    row[1].text = str(b.get("durationMinutes") or "")
    row[2].text = (b.get("laneTitle") or "")[:30]
    row[3].text = b.get("title") or ""

doc.add_page_break()

# ===================== 7. CYCLE EXECUTION =====================
add_header("7. Cycle execution log (30-day rolling)", level=1)
t = make_table(["Date", "Planned", "Completed", "Rate", "Drift", "Streak"])
for d in cycle:
    row = t.add_row().cells
    row[0].text = d.get("date") or ""
    row[1].text = str(d.get("plannedMinutes") or 0)
    row[2].text = str(d.get("completedMinutes") or 0)
    rate = d.get("completionRate")
    row[3].text = f"{rate:.0%}" if isinstance(rate, (int, float)) else ""
    row[4].text = d.get("driftLabel") or ""
    row[5].text = d.get("streakState") or ""

doc.add_page_break()

# ===================== 8. EVALUATION WORKSHEET =====================
add_header("8. Evaluation worksheet", level=1)
doc.add_paragraph(
    "Now you can see the full 5-year plan. Use this section to record your assessment. "
    "Each prompt is followed by space for notes."
)

prompts = [
    (
        "Does the 5-year shape match the real plan?",
        f"{len(fh_blocks)} blocks across {meta.get('range',{}).get('startDayKey')} → {meta.get('range',{}).get('endDayKey')}. "
        "Is the work distribution by year/phase/lane right? Hollow years? Wrong-end-loaded?",
    ),
    (
        "Is the outcome target / success standard captured?",
        "Section 1 shows what Jericho holds as the outcome target. Is this what you actually want?",
    ),
    (
        "Are the 8 lanes the right shape?",
        "Section 4.4 shows blocks-per-lane. Any lane drastically under- or over-resourced relative to its strategic weight?",
    ),
    (
        "Phase distribution (P1/P2/P3) — right shape?",
        "Section 4.1. Is the P1=foundation / P2=operate / P3=terminal weight correct for a 5-year horizon?",
    ),
    (
        "Block-type mix — right shape?",
        "Section 4.5. action/review/audit/validation/readiness/gate — is the proportion right, or is the plan too review-heavy / too action-heavy?",
    ),
    (
        "Year-by-year (Section 5) — what's missing or wrong?",
        "Walk through each year. Flag any block that's mis-titled, mis-timed, or that should not exist. Flag missing work.",
    ),
    (
        "Active cycle (Section 6) — is THIS the right first 30 days?",
        f"{len(cycle_blocks)} blocks form the first executable cycle. Are these the right next moves given everything above?",
    ),
    (
        "What's the biggest gap exposed by this view?",
        "Free-form. What's the single most important thing to fix next?",
    ),
]

for q, hint in prompts:
    p = doc.add_paragraph()
    p.add_run(q).bold = True
    h = doc.add_paragraph()
    hr = h.add_run(hint)
    hr.italic = True
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    for _ in range(6):
        doc.add_paragraph("")

doc.save(OUT)
print(f"WROTE: {OUT}")
print(f"SIZE:  {OUT.stat().st_size} bytes")
