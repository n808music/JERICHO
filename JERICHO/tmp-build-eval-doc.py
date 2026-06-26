#!/usr/bin/env python3
"""Build a Word-openable evaluation doc from the extracted Jericho schedule."""
import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

ROOT = Path("/Users/jamesdotson/vscode/JERICHO/JERICHO")
SRC = ROOT / "tmp-schedule-extract.json"
OUT = ROOT / "tmp-jericho-evaluation.docx"

bundle = json.loads(SRC.read_text())
meta = bundle["meta"]
mp = bundle["masterPlan"]
lanes = bundle["lanes"]
milestones = bundle["milestones"]
blocks = bundle["blocks"]
cycle = bundle["cycleSummary"]

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# Base style
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)


def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x2A, 0x44)
    return h


def add_kv(label, value):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label}: ")
    r1.bold = True
    p.add_run("" if value is None else str(value))
    return p


def fmt_dt(iso):
    if not iso:
        return ""
    try:
        dt = datetime.fromisoformat(iso.replace("Z", "+00:00"))
        return dt.strftime("%a %Y-%m-%d %H:%M UTC")
    except Exception:
        return iso


def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


# ===================== COVER =====================
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("Jericho — Live Plan Evaluation")
tr.font.size = Pt(22)
tr.bold = True

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run(mp.get("title", "(untitled plan)"))
sr.font.size = Pt(14)
sr.italic = True

meta_p = doc.add_paragraph()
meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta_p.add_run(
    f"View date: {meta.get('viewDate','?')}  •  "
    f"Lifecycle: {meta.get('goalLifecycleState','?')}  •  "
    f"Horizon mode: {meta.get('selectedHorizonMode','?')}"
).italic = True

extracted = meta.get("extractedAtISO", "")
stamp = doc.add_paragraph()
stamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
stamp.add_run(f"Extracted: {fmt_dt(extracted)}").font.size = Pt(9)

doc.add_paragraph()

# ===================== 1. PLAN OVERVIEW =====================
add_heading("1. Plan overview", level=1)

add_kv("Plan ID", mp.get("id"))
add_kv("Status", mp.get("status"))
add_kv("Horizon", f"{mp.get('horizonStart')} → {mp.get('horizonEnd')}  ({mp.get('declaredHorizonMonths')} months)")
add_kv("Full-horizon end day key", mp.get("fullHorizonEndDayKey"))
add_kv("Official start", mp.get("officialStartDate"))
add_kv("Schedule applied", fmt_dt(mp.get("scheduleAppliedDate")))
add_kv("Controllability class", mp.get("controllabilityClass"))
add_kv("Terminal target class", mp.get("terminalTargetClass"))
add_kv("Goal architecture", mp.get("goalArchitecture"))
add_kv("Execution model", mp.get("executionModel"))
add_kv("Primary lane", mp.get("primaryLane"))

doc.add_paragraph()
add_heading("North-star outcome", level=2)
doc.add_paragraph(mp.get("northStarOutcome") or "(none)")

add_heading("Core mission", level=2)
doc.add_paragraph(mp.get("coreMission") or "(none)")

add_heading("Outcome target", level=2)
doc.add_paragraph(mp.get("outcomeTarget") or "(none)")

add_heading("Success standard", level=2)
doc.add_paragraph(mp.get("successStandard") or "(none)")

add_heading("Plan summary", level=2)
doc.add_paragraph(mp.get("masterPlanSummary") or "(none)")

# Controllable / external signals
add_heading("Controllable success signals", level=2)
css = mp.get("controllableSuccessSignals") or []
if css:
    for s in css:
        doc.add_paragraph(str(s), style="List Bullet")
else:
    doc.add_paragraph("(none)")

add_heading("Externally-mediated targets", level=2)
emt = mp.get("externallyMediatedTargets") or []
if emt:
    for s in emt:
        doc.add_paragraph(str(s), style="List Bullet")
else:
    doc.add_paragraph("(none)")

add_heading("Non-negotiables", level=2)
nn = mp.get("nonNegotiables") or []
if nn:
    for s in nn:
        doc.add_paragraph(str(s) if not isinstance(s, dict) else json.dumps(s), style="List Bullet")
else:
    doc.add_paragraph("(none)")

add_heading("Anchors", level=2)
anchors = mp.get("anchors") or []
if anchors:
    for a in anchors:
        doc.add_paragraph(json.dumps(a) if isinstance(a, dict) else str(a), style="List Bullet")
else:
    doc.add_paragraph("(none)")

# Quality
add_heading("Quality gate (applied)", level=2)
add_kv("Policy", meta.get("qualityPolicyIdApplied"))
add_kv("Score", meta.get("qualityScoreApplied"))
comp = meta.get("qualityScoreAppliedByComponent") or {}
if comp:
    for k, v in comp.items():
        add_kv(f"  • {k}", v)

doc.add_page_break()

# ===================== 2. LANES =====================
add_heading("2. Lanes", level=1)
doc.add_paragraph(f"{len(lanes)} lanes configured. Each lane is a parallel workstream.")

t = doc.add_table(rows=1, cols=5)
t.style = "Light List Accent 1"
hdr = t.rows[0].cells
for i, label in enumerate(["Lane", "Domain", "Role", "Activation", "Stage"]):
    hdr[i].text = label
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
    shade_cell(hdr[i], "1F2A44")
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for lane_id, lane in lanes.items():
    row = t.add_row().cells
    row[0].text = lane.get("title") or "(untitled)"
    row[1].text = lane.get("domain") or ""
    row[2].text = lane.get("role") or ""
    row[3].text = lane.get("activationState") or ""
    row[4].text = lane.get("assessedStage") or ""

# Lane descriptions
doc.add_paragraph()
for lane_id, lane in lanes.items():
    add_heading(lane.get("title") or "(untitled)", level=3)
    desc = lane.get("userDescription") or "(no description)"
    doc.add_paragraph(desc)
    add_kv("Domain / role", f"{lane.get('domain')} / {lane.get('role')}")
    add_kv("Activation", lane.get("activationState"))
    add_kv("Assessed stage", lane.get("assessedStage"))
    add_kv("Priority score", lane.get("priorityScore"))
    add_kv("Milestone count", len(lane.get("milestoneIds") or []))

doc.add_page_break()

# ===================== 3. MILESTONES =====================
add_heading("3. Milestones (sorted by target date)", level=1)
doc.add_paragraph(f"{len(milestones)} milestones across {len(lanes)} lanes.")

mt = doc.add_table(rows=1, cols=5)
mt.style = "Light List Accent 1"
mhdr = mt.rows[0].cells
for i, label in enumerate(["Target date", "Lane", "Milestone", "Type", "Flex"]):
    mhdr[i].text = label
    for p in mhdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
    shade_cell(mhdr[i], "1F2A44")
    for p in mhdr[i].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for m in milestones:
    row = mt.add_row().cells
    row[0].text = m.get("targetDate") or ""
    row[1].text = (m.get("laneTitle") or "")[:30]
    row[2].text = m.get("title") or ""
    row[3].text = m.get("milestoneType") or ""
    row[4].text = m.get("flex") or ""

# Miss-consequence detail
doc.add_paragraph()
add_heading("Milestone miss consequences", level=2)
for m in milestones:
    mc = (m.get("missConsequence") or "").strip()
    if not mc:
        continue
    p = doc.add_paragraph()
    p.add_run(f"{m.get('targetDate','')} — {m.get('title','')}: ").bold = True
    p.add_run(mc)

doc.add_page_break()

# ===================== 4. SCHEDULED BLOCKS =====================
add_heading("4. Scheduled blocks (chronological)", level=1)
total_min = sum(b.get("durationMinutes", 0) or 0 for b in blocks)
doc.add_paragraph(
    f"{len(blocks)} blocks scheduled, {total_min} minutes total ({total_min/60:.1f} hours). "
    f"Range: {blocks[0]['dayKey'] if blocks else '—'} → {blocks[-1]['dayKey'] if blocks else '—'}."
)

bt = doc.add_table(rows=1, cols=5)
bt.style = "Light List Accent 1"
bhdr = bt.rows[0].cells
for i, label in enumerate(["Start (UTC)", "Min", "Lane", "Title", "Expected output"]):
    bhdr[i].text = label
    for p in bhdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
    shade_cell(bhdr[i], "1F2A44")
    for p in bhdr[i].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for b in blocks:
    row = bt.add_row().cells
    row[0].text = fmt_dt(b.get("startISO"))
    row[1].text = str(b.get("durationMinutes") or "")
    lane_title = lanes.get(b.get("laneId") or "", {}).get("title") or "—"
    row[2].text = lane_title[:24]
    row[3].text = b.get("title") or ""
    row[4].text = (b.get("expectedOutput") or "")[:120]

# Per-block detail
doc.add_paragraph()
add_heading("Per-block detail", level=2)
for i, b in enumerate(blocks, 1):
    add_heading(f"Block {i}: {b.get('title','(untitled)')}", level=3)
    add_kv("Start", fmt_dt(b.get("startISO")))
    add_kv("End", fmt_dt(b.get("endISO")))
    add_kv("Duration", f"{b.get('durationMinutes')} min")
    add_kv("Domain", b.get("domain"))
    add_kv("Block type", b.get("blockType"))
    add_kv("Milestone type", b.get("milestoneType"))
    add_kv("Flex", b.get("flex"))
    add_kv("Lane", lanes.get(b.get("laneId") or "", {}).get("title") or "—")
    add_kv("Status", b.get("status"))
    add_kv("Source", b.get("source"))
    add_kv("Placement basis", b.get("placementBasis"))
    if b.get("expectedOutput"):
        add_kv("Expected output", b.get("expectedOutput"))
    if b.get("missConsequence"):
        add_kv("Miss consequence", b.get("missConsequence"))
    if b.get("sourceQuestion"):
        add_kv("Source question", b.get("sourceQuestion"))
    if b.get("lineageTitle"):
        add_kv("Lineage", b.get("lineageTitle"))

doc.add_page_break()

# ===================== 5. CYCLE SUMMARY =====================
add_heading("5. Cycle (30-day rolling window)", level=1)
doc.add_paragraph("Day-by-day planned vs completed minutes for the active cycle.")

ct = doc.add_table(rows=1, cols=6)
ct.style = "Light List Accent 1"
chdr = ct.rows[0].cells
for i, label in enumerate(["Date", "Planned", "Completed", "Rate", "Drift", "Streak"]):
    chdr[i].text = label
    for p in chdr[i].paragraphs:
        for r in p.runs:
            r.bold = True
    shade_cell(chdr[i], "1F2A44")
    for p in chdr[i].paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for d in cycle:
    row = ct.add_row().cells
    row[0].text = d.get("date") or ""
    row[1].text = str(d.get("plannedMinutes") or 0)
    row[2].text = str(d.get("completedMinutes") or 0)
    rate = d.get("completionRate")
    row[3].text = f"{rate:.0%}" if isinstance(rate, (int, float)) else ""
    row[4].text = d.get("driftLabel") or ""
    row[5].text = d.get("streakState") or ""

doc.add_page_break()

# ===================== 6. EVALUATION =====================
add_heading("6. Evaluation worksheet", level=1)
doc.add_paragraph(
    "Use this section to record your assessment of the plan above. "
    "Each prompt is followed by blank lines for handwritten or typed notes."
)

prompts = [
    ("Does the north-star outcome match what you actually want?",
     "Is the outcome target falsifiable? Is the success standard concrete enough to know when you've hit it?"),
    ("Are the right lanes activated?",
     "Should any inactive lane be activated? Should any active lane be paused? Is anything missing?"),
    ("Are the milestone dates realistic?",
     "Look at the 39 milestones across 5 years. Flag any whose target date feels wrong (too aggressive, too lax, mis-ordered)."),
    ("Are the next 24 scheduled blocks the right next moves?",
     "For the May-1 → June-9 window: does the sequence build toward the right milestones? Anything missing? Anything filler?"),
    ("Is the time load realistic?",
     "22.3 hours over 40 days ≈ 33 min/day. Is that the right intensity for this horizon?"),
    ("Is the FOCUS / CREATION / RESOURCES mix right?",
     "Current split: FOCUS 9, RESOURCES 8, CREATION 7. Should it be tilted differently?"),
    ("Are the four no-lane blocks a problem?",
     "4 of 24 blocks have no lane assignment. Should they be assigned, or are they legitimately cross-cutting?"),
    ("Capital stack has only 1 block. Is that under-served?",
     "Compare to album / product / runway / brand-ops (4 each) and media (3). One block for capital — intentional or a gap?"),
    ("Does the plan-quality score (20, BALANCED policy) feel right?",
     "If you'd describe this plan in one word, is 'BALANCED' the right word? Why is the score only 20?"),
    ("What would you change?",
     "Free-form: anything Jericho got wrong, anything it surfaced that you hadn't noticed."),
]

for q, hint in prompts:
    p = doc.add_paragraph()
    p.add_run(q).bold = True
    h = doc.add_paragraph()
    hr = h.add_run(hint)
    hr.italic = True
    hr.font.size = Pt(9)
    hr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    for _ in range(5):
        doc.add_paragraph("")

doc.save(OUT)
print(f"WROTE: {OUT}")
print(f"SIZE:  {OUT.stat().st_size} bytes")
